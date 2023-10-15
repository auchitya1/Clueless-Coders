from bs4 import BeautifulSoup
from docx import Document

# Function to convert HTML to DOCX
#def html_to_docx(input_html, output_docx):
    # Load HTML content using BeautifulSoup
with open(input_html, 'r', encoding='utf-8') as file:
    html_content = file.read()
soup = BeautifulSoup(html_content, 'lxml')

    # Create a new DOCX document
doc = Document()
    
    # Iterate through HTML elements and convert to DOCX
for element in soup.find_all():
    if element.name == 'p':
        doc.add_paragraph(element.get_text())
    elif element.name == 'h1':
        doc.add_heading(element.get_text(), level=1)
    elif element.name == 'h2':
        doc.add_heading(element.get_text(), level=2)
        # Handle other HTML elements as needed, e.g., tables, lists, hyperlinks
        

    # Save the DOCX file
doc.save(output_docx)

# Example usage
#html_to_docx('sample.html', 'output.docx')

# Locate and extract CSS styles from the HTML
soup = BeautifulSoup(html_content, 'lxml')
style_tags = soup.find_all('style')
css_styles = "\n".join([style_tag.string for style_tag in style_tags])
for paragraph in doc.paragraphs:
    if paragraph.style.name == 'Normal':
        paragraph.style.font.name = 'Arial'
        paragraph.style.font.size = Pt(14)
        paragraph.style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
# Handle text alignment
if 'text-align' in css_styles:
    alignment = css_styles['text-align']
    if alignment == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Handle other alignment options (left, right, justify, etc.)
# Extract color styles from HTML
text_color = rule.get('color')  # Replace 'rule' with the appropriate selector
background_color = rule.get('background-color')

from docx.shared import RGBColor

for paragraph in doc.paragraphs:
    if paragraph.style.name == 'Normal':
        if text_color:
            # Apply text color to runs within the paragraph
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(*parse_color(text_color))
        if background_color:
            # Apply background color to paragraph shading
            paragraph.paragraph_format.shading.background_color = RGBColor(*parse_color(background_color))

def parse_color(color_value):
    color_value = color_value.strip()
    if color_value.startswith("rgba("):
        parts = color_value[5:-1].split(',')
        r, g, b, a = int(parts[0]), int(parts[1]), int(parts[2]), float(parts[3])
        return (r, g, b, a)
    

# Extract margin and spacing styles from HTML
margin = rule.get('margin')  # Replace 'rule' with the appropriate selector
padding = rule.get('padding')
line_height = rule.get('line-height')
text_indent = rule.get('text-indent')

for paragraph in doc.paragraphs:
    if paragraph.style.name == 'Normal':
        if margin:
            # Apply margin size to paragraph
            paragraph.paragraph_format.space_before = Inches(float(margin))
            paragraph.paragraph_format.space_after = Inches(float(margin))
        if padding:
            # Apply padding size to paragraph (not directly supported in DOCX)
            # You may need to adjust spacing within paragraphs manually
        if line_height:
            # Apply line spacing to paragraph
            paragraph.paragraph_format.line_spacing = Pt(float(line_height))
        if text_indent:
            # Apply text indentation to paragraph
            paragraph.paragraph_format.first_line_indent = Pt(float(text_indent))

for ul in soup.find_all('ul'):
    # Create a bullet list in the DOCX document
    doc.add_paragraph('', style='List Bullet')  # Apply the bullet style
    for li in ul.find_all('li'):
        # Add list items to the bullet list
        doc.add_paragraph(li.get_text(), style='List Bullet')
        # Handle nested lists within <li> elements (recursively)
        for nested_ul in li.find_all('ul'):
            doc.add_paragraph('', style='List Bullet')
            for nested_li in nested_ul.find_all('li'):
                doc.add_paragraph(nested_li.get_text(), style='List Bullet')

from docx.enum.text import WD_ALIGN_PARAGRAPH

for table in soup.find_all('table'):
    # Create a table in the DOCX document
    doc.add_table(rows=2, cols=2).style = 'Table Grid'  # Apply a table style
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Apply table alignment

    # Fill in table cells
    row_index = 0
    for row in table.find_all('tr'):
        col_index = 0
        for cell in row.find_all('td'):
            # Access table cells and apply styles
            cell_paragraph = doc.tables[-1].cell(row_index, col_index).paragraphs[0]
            cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Cell alignment
            cell_paragraph.paragraph_format.space_after = Pt(0)  # Cell spacing
            cell_paragraph.paragraph_format.space_before = Pt(0)
            cell_paragraph.paragraph_format.line_spacing = Pt(0)
            cell_paragraph.paragraph_format.left_indent = Pt(0)  # Text indentation

            # Set cell border style
            cell_paragraph.paragraph_format.border_around(1, Pt(1), RGBColor(0, 0, 0))  # Replace with your border styles
            cell_paragraph.runs[0].bold = True  # Example: make text bold

            # Fill cell with content
            cell_paragraph.add_run(cell.get_text())

            col_index += 1
        row_index += 1

for paragraph in soup.find_all('p'):
    for special_char in paragraph.find_all(True, {'class': 'special-char'}):  # Identify special characters with a specific class
        entity = special_char.string  # This will contain the HTML entity, e.g., '&copy;'

import html

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.text = html.unescape(run.text)  # Convert HTML entities to Unicode characters


    